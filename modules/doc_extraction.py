
import logging
import os
from pathlib import Path
from typing import Optional

import fitz               
import docx                
import pytesseract
from PIL import Image

from modules.config import OCR_LANG

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    
    full_text = []
    try:
        with fitz.open(file_path) as pdf_doc:
            total_pages = pdf_doc.page_count
            logger.info(f"Processing PDF: {os.path.basename(file_path)} ({total_pages} pages)")

            for page_num, page in enumerate(pdf_doc, start=1):
                digital_text = page.get_text("text").strip()

                if len(digital_text) >= 50:
                    full_text.append(digital_text)
                else:
                    # Scanned page — render to image then OCR
                    logger.debug(
                        f"  Page {page_num}/{total_pages}: "
                        f"only {len(digital_text)} chars — falling back to OCR"
                    )
                    mat = fitz.Matrix(200 / 72, 200 / 72)   # 200 DPI
                    pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img, lang=OCR_LANG)
                    full_text.append(ocr_text.strip())

        result = "\n\n--- Page Break ---\n\n".join(full_text)
        logger.info(f"PDF extraction complete: {len(result):,} characters")
        return result

    except fitz.FileDataError as e:
        logger.error(f"Corrupt or password-protected PDF '{file_path}': {e}")
        raise RuntimeError(f"Cannot read PDF (corrupt or encrypted): {e}") from e
    except Exception as e:
        logger.error(f"PDF extraction failed for '{file_path}': {e}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    
    try:
        doc = docx.Document(file_path)
        sections = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                sections.append(text)

        # Tables — very common in government legal documents
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    sections.append(row_text)

        result = "\n\n".join(sections)
        logger.info(
            f"DOCX extraction complete: {len(sections)} blocks, "
            f"{len(result):,} characters from '{os.path.basename(file_path)}'"
        )
        return result

    except Exception as e:
        logger.error(f"DOCX extraction failed for '{file_path}': {e}")
        raise

def extract_text_from_txt(file_path: str) -> str:

    try:
        text = Path(file_path).read_text(encoding="utf-8", errors="replace")
        logger.info(f"TXT read: {len(text):,} characters from '{os.path.basename(file_path)}'")
        return text.strip()
    except Exception as e:
        logger.error(f"TXT read failed for '{file_path}': {e}")
        raise

def extract_text_from_image(file_path: str) -> str:
    
    try:
        image = Image.open(file_path).convert("L")   # grayscale
        custom_config = r"--oem 3 --psm 6"
        text = pytesseract.image_to_string(image, lang=OCR_LANG, config=custom_config)
        result = text.strip()
        logger.info(f"Image OCR: {len(result):,} chars from '{os.path.basename(file_path)}'")
        return result
    except Exception as e:
        logger.error(f"Image OCR failed for '{file_path}': {e}")
        raise

_EXTRACTORS = {
    ".pdf":  extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".doc":  extract_text_from_docx,
    ".txt":  extract_text_from_txt,
    ".png":  extract_text_from_image,
    ".jpg":  extract_text_from_image,
    ".jpeg": extract_text_from_image,
}

def process_uploaded_file(file_path: str) -> str:

    ext = Path(file_path).suffix.lower()
    extractor = _EXTRACTORS.get(ext)

    if extractor is None:
        supported = ", ".join(_EXTRACTORS.keys())
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {supported}"
        )

    logger.info(f"Extracting '{os.path.basename(file_path)}' using {extractor.__name__}")
    return extractor(file_path)

def get_supported_extensions() -> list:
    """Returns the list of supported file extensions."""
    return list(_EXTRACTORS.keys())
